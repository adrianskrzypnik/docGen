from docx import Document
import sys
import json

def create_mode1_document(data):
    doc = Document()
    doc.add_heading('Tryb 1 - Informacje o kliencie i wykonawcy', level=0)
    doc.add_paragraph(f"Klient: {data.get('client', 'Brak danych')}")
    doc.add_paragraph(f"Wykonawca: {data.get('contractor', 'Brak danych')}")
    doc.add_heading('Daty szkoleń', level=1)
    for key, value in data.get('trainingDates', {}).items():
        doc.add_paragraph(f"{key}: {value}")
    doc.add_heading('Daty podpisu', level=1)
    for key, value in data.get('signatureDates', {}).items():
        doc.add_paragraph(f"{key}: {value}")
    return doc

def create_document(json_data):
    try:
        data_dict = json.loads(json_data)
    except json.JSONDecodeError as e:
        print(f"Error decoding main JSON: {e}", file=sys.stderr)
        return 'Error decoding main JSON'
    if 'data' not in data_dict:
        print("Error: Missing 'data' key", file=sys.stderr)
        return 'Error: Missing "data" key'
    try:
        inner_data = json.loads(data_dict['data'])
    except json.JSONDecodeError as e:
        print(f"Error decoding 'data' JSON: {e}", file=sys.stderr)
        return 'Error decoding "data" JSON'
    mode = inner_data.get('mode')
    content = inner_data.get('data', {})
    mode_functions = {'mode1': create_mode1_document}
    if mode not in mode_functions:
        print(f"Error: Unsupported mode '{mode}'", file=sys.stderr)
        return f'Error: Unsupported mode "{mode}"'
    try:
        doc = mode_functions[mode](content)
        doc.save('output.docx')
        return 'Document created successfully!'
    except Exception as e:
        print(f"Error saving document: {e}", file=sys.stderr)
        return 'Error saving document'

if __name__ == "__main__":
    try:
        result = create_document(sys.argv[1])
        print(json.dumps({'status': result}))
    except Exception as e:
        print(json.dumps({'status': 'error', 'message': str(e)}))
